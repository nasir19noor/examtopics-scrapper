#!/usr/bin/env python3
"""
make_docx.py — build a Word study doc from the scraped questions + authored
answers, matching the format of output/sap.docx.

Standalone helper; it does not touch the other scripts.

Per question, matching sap.docx:
  * question stem   — plain black, first line prefixed "N. "
  * A/B/C/D/E lines — plain black, except the correct choice(s): RED + bold
  * "Explanation : " and "Answer: X" — GREEN (Answer bold)
  * the authored explanation body — GREEN

Sources:
  * questions (stem + choices)  : output/aip-c01/<n>.txt  (the scrape)
  * correct letter + explanation: output/answers.json     (authored)

sap.docx is used as the template so the new document inherits its exact
'normal1' style (Calibri 10pt) and page setup.

Usage:
    python make_docx.py
    python make_docx.py --template output/sap.docx --src output/aip-c01 \
                        --answers output/answers.json --out output/aws-aip-c01.docx
"""

from __future__ import annotations

import argparse
import glob
import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)      # correct answer         (matches sap.docx)
GREEN = RGBColor(0x18, 0x80, 0x38)    # explanation section    (matches sap.docx)

_CHOICE_RE = re.compile(r"^([A-E])\.\s")


def parse_question(path: str) -> tuple[list[str], list[tuple[str, str]]]:
    """Return (stem_paragraphs, [(letter, choice_text), …]) from a scrape file."""
    text = path if "\n" in path else Path(path).read_text(encoding="utf-8")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Drop the Exam:/Question #: header and everything from the answer onward.
    text = re.sub(r"^Exam:.*\n", "", text, count=1)
    text = re.sub(r"^Question #:.*\n", "", text, count=1)
    text = re.split(r"\nCorrect answer:", text, maxsplit=1)[0]

    stem: list[str] = []
    choices: list[tuple[str, str]] = []
    for line in text.split("\n"):
        line = line.rstrip()
        if not line.strip():
            continue
        m = _CHOICE_RE.match(line)
        if m:
            choices.append((m.group(1), line))
        elif choices:                 # a wrapped continuation of the last choice
            letter, prev = choices[-1]
            choices[-1] = (letter, prev + " " + line.strip())
        else:
            stem.append(line.strip())
    return stem, choices


def clear_body(doc: Document) -> None:
    """Remove every existing paragraph so we reuse only the template's styles."""
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)


def add_line(doc: Document, text: str, color=None, bold=False):
    """Add a paragraph, rendering **markdown bold** as real bold runs."""
    p = doc.add_paragraph(style="normal1")
    for seg in re.split(r"(\*\*.+?\*\*)", text):
        if not seg:
            continue
        emph = seg.startswith("**") and seg.endswith("**") and len(seg) > 4
        run = p.add_run(seg[2:-2] if emph else seg)
        if color is not None:
            run.font.color.rgb = color
        run.bold = bold or emph
    return p


def build(template: str, src_dir: str, answers_path: str, out_path: str) -> int:
    answers = json.loads(Path(answers_path).read_text(encoding="utf-8"))
    files = {int(os.path.basename(f)[:-4]): f
             for f in glob.glob(os.path.join(src_dir, "*.txt"))
             if os.path.basename(f)[:-4].isdigit()}

    doc = Document(template)
    clear_body(doc)

    missing = []
    for i, n in enumerate(sorted(files)):
        stem, choices = parse_question(files[n])
        entry = answers.get(str(n)) or {}
        correct = set(entry.get("answer", "") or "")
        explanation = entry.get("explanation", "")
        if not correct or not explanation:
            missing.append(n)

        # Stem — first line carries the "N. " number.
        for j, para in enumerate(stem):
            add_line(doc, f"{n}. {para}" if j == 0 else para)

        # Choices — correct one(s) in red bold.
        for letter, ctext in choices:
            if letter in correct:
                add_line(doc, ctext, color=RED, bold=True)
            else:
                add_line(doc, ctext)

        # Explanation block — green.
        add_line(doc, "Explanation : ", color=GREEN)
        add_line(doc, f"Answer: {entry.get('answer', '?')}", color=GREEN, bold=True)
        for para in explanation.split("\n"):
            if para.strip():
                add_line(doc, para.strip(), color=GREEN)

        if i != len(files) - 1:
            doc.add_paragraph(style="normal1")     # blank line between questions

    doc.save(out_path)
    print(f"wrote {len(files)} questions -> {out_path}")
    if missing:
        print(f"WARNING: {len(missing)} questions have no authored answer yet: {missing}")
    return 0


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Build the Word study doc.")
    ap.add_argument("--template", default="output/sap.docx")
    ap.add_argument("--src", default="output/aip-c01")
    ap.add_argument("--answers", default="output/answers.json")
    ap.add_argument("--out", default="output/aws-aip-c01.docx")
    args = ap.parse_args(argv)
    return build(args.template, args.src, args.answers, args.out)


if __name__ == "__main__":
    sys.exit(main())
